import csv
import sys 
import os
import re
import nltk 
import PyPDF2 
import xtract
import textract
from tkinter import *

from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from docx import *
from docx.shared import Inches

nltk.download('averaged_perceptron_tagger')
nltk.download('stopwords')
nltk.download('punkt')

document = Document()

# def run():
#     entry_value = entry.get()
#     print(entry_value)

screen = Tk()

screen.title("Keyword Finder")

open_text = Label(text = "Enter your PDF Location \n 'Example:Users/Name/Location/book.pdf'", fg = "black", bg = "white")
open_text.pack()

top = Frame(screen)
top.pack(side = TOP)

entry = Entry(top)
entry.pack()
# entry_value = 'engines_of_creation.pdf'
key_value = ''

# This function will extract and return the pdf file text content.
def run(filePath=entry.get()):

    directory_with_files_of_interest = "../../../"
    file_to_convert_to_txt = entry.get()
    converted_filename = "text.txt"
    #scroll over so you don't miss cut off text here
    os.system("python ./pdf2text.py -o %s %s/%s" %(converted_filename, directory_with_files_of_interest, file_to_convert_to_txt))

    #take a look at the contents
    file = open("%s" %(converted_filename), "r+")
    for line in file:
        tokens      = word_tokenize(line)
        text        = nltk.Text(tokens)
        words       = [w.lower() for w in tokens]
        keywords    = sorted(set(words))
        texts       = nltk.Text(word for word in nltk.pos_tag(tokens))
        finalText = ' '.join(keywords)
        keys = str(keywords)
        # def search(text):
            # text.concordance('so')
            # text.concordance('and')
            # text.concordance('or')
            # text.concordance('but')
            # text.concordance('nor')
            # text.concordance('for')
            # text.concordance('yet')
            # text.concordance('as')
            # text.concordance('whether')
            # text.concordance('while')
            # text.concordance('what')
            # text.concordance('when')
            # text.concordance('whenever')
            # text.concordance('wherever')
            # text.concordance('even if')
            # text.concordance('even though')
            # text.concordance('once')
            # text.concordance('since')
            # text.concordance('so that')
            # text.concordance('though')
            # text.concordance('till')
            # text.concordance('unless')
            # text.concordance('until')
            # text.concordance('as if')
            # text.concordance('as long as')
            # text.concordance('because')
            # text.concordance('before')
            # text.concordance('although')
            # text.similar('area')
            # text.similar('book')
            # text.similar('business')
            # text.similar('case')
            # text.similar('child')
            # text.similar('country')
            # text.similar('day')
            # text.similar('eye')
            # text.similar('fact')
            # text.similar('group')
            # text.similar('hand')
            # text.similar('home')
            # text.similar('job')
            # text.similar('life')
            # text.similar('woman')
            # text.similar('money')
            # text.similar('month')
            # text.similar('number')
            # text.similar('part')
            # text.similar('problem')
            # text.similar('system')
            # text.similar('water')
            # text.similar('break')
            # text.similar('dead')
            # text.similar('advance')
            # text.similar('pass')
            # text.similar('charge')
            # text.similar('be')
            # text.similar('have')
            # text.similar('do')
            # text.similar('say')
            # text.similar('get')
            # text.similar('make')
            # text.similar('go')
            # text.similar('know')
            # text.similar('take')
            # text.similar('see')
            # text.similar('come')
            # text.similar('think')
            # text.similar('look')
            # text.similar('want')
            # text.similar('give')
            # text.similar('use')
            # text.similar('find')
            # text.similar('tell')
            # text.similar('seem')
            # text.similar('feel')
            # text.similar('try')
            # text.similar('leave')
            # text.similar('call')
            # text.similar('order')
            # text.similar('move')
            # text.similar('solid')
            # text.similar('big')

        # written = search(text)  
        # finished = str(written)
        # results = Text(bottom, width=30, height=1)
        # results.pack()
        # results.insert(0.0, finished)
        # print(tokens)   
        # print(text)     
        # print(words)    
        # print(keywords) 
        # print(texts)    
        # print(finalText)
        
        # for l in finished:
        #     if l is None or l is 'no matches':
        #         del l
        #     else: 
        #         token = word_tokenize(l)
        #         texts = nltk.Text(token)
        #         word = [w.lower() for w in token]
        #         keyword = sorted(set(word))
        #         break 
        otherresults = Listbox(screen, width=30, height=1, yscrollcommand=scroll.set)
        for line in range(1000):
            otherresults.insert(END, finalText)
        otherresults.pack()
        # otherresults.insert(0.0, finalText)
        scroll.config(command=otherresults.yview)
        # otherresults.config(yscrollcommand=scroll.set)
        # final = open("demo.txt", "w+")
        # final.write(keys)
        # final.close()
    file.close()
    print('Done!')
    return

click_me = Button(top, text="Submit", command = run)
click_me.place(x=215, y=60)
click_me.pack()

bottom = Frame(screen)

scroll = Scrollbar(screen)
scroll.pack(side=RIGHT, fill=Y)


bottom.pack()

screen.mainloop()

#Dataframe with unique keywords to avoid repetition in rows
#df = pd.DataFrame(list(set(keywords)),columns=['keywords'])

# results_value = keywords
# results.insert(INSERT, results_value)


# table = document.add_table(rows=1, cols=3)

# with open('demo.csv', 'w') as csvFile:
#     writer = csv.writer(csvFile)
#     writer.writerows(keywords)
# csvFile.close()

# document.add_paragraph(pdfText, style='List Bullet')
# document.save('demo.docx')

# lastThing = open("texts.docx", "w+")
# path = './texts.docx'
# direct = os.listdir(path)
# for i in direct:
#     document = Document()
#     myfile = open(path+i).write(finished)
#     myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile)
#     document.add_paragraph(myfile)
#     document.save('./'+i+'.docx')

# final = open("demo.txt", "w+")
# final.write(finished)
# final.close()

# lastThing.close()


